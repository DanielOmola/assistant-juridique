# utils/utils_tampon.py
import yaml
from datetime import datetime
import os

# ============================================
# CHARGEMENT / SAUVEGARDE CONFIG
# ============================================

def charger_config_tampon(chemin_config=None):
    """
    Charge la configuration du tampon depuis un fichier YAML
    """
    if chemin_config is None:
        chemin_config = "/content/drive/MyDrive/assistant-juridique/config/tampon_config.yaml"
    
    # Configuration par défaut
    config_defaut = {
        "avocat": {
            "nom": "[NOM]",
            "prenom": "[PRÉNOM]",
            "fonction": "Avocat",
            "barreau": "[BARREAU]",
            "siret": "[SIRET]",
            "telephone": "[TÉLÉPHONE]",
            "email": "[EMAIL]",
            "adresse": "[ADRESSE]"
        },
        "options": {
            "type_tampon": "officiel",
            "position": "debut",
            "date_format": "%d/%m/%Y"
        }
    }
    
    try:
        with open(chemin_config, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
            if config:
                return config
            return config_defaut
    except FileNotFoundError:
        # Créer le dossier config si besoin
        os.makedirs(os.path.dirname(chemin_config), exist_ok=True)
        # Sauvegarder la config par défaut
        with open(chemin_config, 'w', encoding='utf-8') as f:
            yaml.dump(config_defaut, f, allow_unicode=True, default_flow_style=False)
        return config_defaut
    except Exception as e:
        print(f"⚠️ Erreur chargement config: {e}")
        return config_defaut

def sauvegarder_config_tampon(config, chemin_config=None):
    """
    Sauvegarde la configuration du tampon
    """
    if chemin_config is None:
        chemin_config = "/content/drive/MyDrive/assistant-juridique/config/tampon_config.yaml"
    
    os.makedirs(os.path.dirname(chemin_config), exist_ok=True)
    
    with open(chemin_config, 'w', encoding='utf-8') as f:
        yaml.dump(config, f, allow_unicode=True, default_flow_style=False)
    
    return True

# ============================================
# GÉNÉRATION DU TAMPON
# ============================================

def generer_tampon(config=None, type_tampon=None):
    """
    Génère un tampon minimaliste avec toutes les mentions obligatoires
    """
    if config is None:
        config = charger_config_tampon()
    
    avocat = config["avocat"]
    options = config["options"]
    
    type_tampon = type_tampon or options.get("type_tampon", "officiel")
    date_str = datetime.now().strftime(options.get("date_format", "%d/%m/%Y"))
    
    # Mentions obligatoires (minimales mais complètes)
    mentions = f"""{avocat['prenom']} {avocat['nom']}
{avocat['fonction']}
Barreau de {avocat['barreau']}
SIRET: {avocat['siret']}
Tél: {avocat['telephone']}
{avocat['email']}
{avocat['adresse']}"""
    
    # Version officielle (mentions complètes)
    tampon_officiel = f"""{mentions}
{date_str}"""
    
    # Version confidentielle (mentions allégées)
    tampon_confidentiel = f"""🔒 CONFIDENTIEL
{avocat['prenom']} {avocat['nom']}
{avocat['fonction']} - Barreau de {avocat['barreau']}
{date_str}"""
    
    # Version reçu
    tampon_recu = f"""✅ REÇU LE {date_str}
{avocat['prenom']} {avocat['nom']}
{avocat['fonction']}"""
    
    # Version brouillon
    tampon_brouillon = f"""📝 BROUILLON
{avocat['prenom']} {avocat['nom']}
{date_str}"""
    
    tampons = {
        "officiel": tampon_officiel,
        "confidentiel": tampon_confidentiel,
        "recu": tampon_recu,
        "brouillon": tampon_brouillon
    }
    
    return tampons.get(type_tampon, tampon_officiel)

def appliquer_tampon(texte, config=None, type_tampon=None, position=None):
    """
    Applique le tampon à un texte brut
    """
    if config is None:
        config = charger_config_tampon()
    
    type_tampon = type_tampon or config["options"].get("type_tampon", "officiel")
    position = position or config["options"].get("position", "debut")
    
    tampon = generer_tampon(config, type_tampon)
    
    if position == "debut":
        return tampon + "\n\n" + texte
    elif position == "fin":
        return texte + "\n\n" + tampon
    else:
        return tampon + "\n\n" + texte + "\n\n" + tampon

# ============================================
# TAMPONNAGE DES FICHIERS (PDF, DOCX, TXT)
# ============================================

def appliquer_tampon_fichier(chemin_entree, chemin_sortie=None, config=None, 
                              type_tampon=None, position=None):
    """
    Applique un tampon à un fichier (conserve le format)
    """
    if config is None:
        config = charger_config_tampon()
    
    type_tampon = type_tampon or config["options"].get("type_tampon", "officiel")
    position = position or config["options"].get("position", "debut")
    
    tampon_texte = generer_tampon(config, type_tampon)
    
    if chemin_sortie is None:
        base, ext = os.path.splitext(chemin_entree)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        chemin_sortie = f"{base}_tamponne_{timestamp}{ext}"
    
    extension = os.path.splitext(chemin_entree)[1].lower()
    
    try:
        # Texte brut
        if extension in ['.txt', '.md']:
            with open(chemin_entree, 'r', encoding='utf-8') as f:
                texte = f.read()
            
            if position == "debut":
                texte_tamponne = tampon_texte + "\n\n" + texte
            elif position == "fin":
                texte_tamponne = texte + "\n\n" + tampon_texte
            else:
                texte_tamponne = tampon_texte + "\n\n" + texte + "\n\n" + tampon_texte
            
            with open(chemin_sortie, 'w', encoding='utf-8') as f:
                f.write(texte_tamponne)
        
        # Word
        elif extension == '.docx':
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document(chemin_entree)
            section = doc.sections[0]
            
            if position in ["debut", "debut_et_fin"]:
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.text = tampon_texte
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            
            if position in ["fin", "debut_et_fin"]:
                footer = section.footer
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                paragraph.text = tampon_texte
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            
            doc.save(chemin_sortie)
        
        # PDF
        elif extension == '.pdf':
            from PyPDF2 import PdfReader, PdfWriter
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            import io
            
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=letter)
            c.setFont("Helvetica", 8)
            
            y_position = 750 if position in ["debut", "debut_et_fin"] else 50
            
            for i, ligne in enumerate(tampon_texte.split('\n')):
                c.drawString(50, y_position - (i * 10), ligne[:80])
            
            c.save()
            packet.seek(0)
            
            tampon_pdf = PdfReader(packet)
            pdf_original = PdfReader(chemin_entree)
            pdf_sortie = PdfWriter()
            
            for i, page in enumerate(pdf_original.pages):
                if position in ["debut", "debut_et_fin"] and i == 0:
                    page.merge_page(tampon_pdf.pages[0])
                if position in ["fin", "debut_et_fin"] and i == len(pdf_original.pages) - 1:
                    page.merge_page(tampon_pdf.pages[0])
                pdf_sortie.add_page(page)
            
            with open(chemin_sortie, 'wb') as f:
                pdf_sortie.write(f)
        
        else:
            return None, f"Format {extension} non supporté"
        
        return chemin_sortie, None
        
    except Exception as e:
        return None, str(e)

print("✅ Module utils_tampon chargé avec succès")